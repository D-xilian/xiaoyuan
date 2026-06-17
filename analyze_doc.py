from docx import Document

doc = Document(r'd:\校园活动\基于VUE+Flask架构的校园活动发布与管理系统_论文.docx')

print(f"总段落数: {len(doc.paragraphs)}")
print(f"总表格数: {len(doc.tables)}")
print()

for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    text = p.text[:200] if p.text else ''
    if text.strip():
        print(f"[{i}] Style={style} | {text}")
