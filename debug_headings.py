# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document(r'd:\校园活动\基于VUE+Flask架构的校园活动发布与管理系统_论文_修改版.docx')

# 检查标题段落的实际内容
print('=== 检查所有标题段落 ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t and (t.startswith('第') or t.startswith('3.') or t.startswith('4') or '轻量级' in t or '研究问题' in t):
        print(f'[{i}] style={p.style.name} | runs={len(p.runs)} | text="{t}"')
        # Show the actual XML structure briefly
        for j, run in enumerate(p.runs):
            print(f'  run[{j}]: font={run.font.name}, text="{run.text}"')

print()
print('=== 检查第4章 ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if '结论' in t and (t.startswith('4') or t.startswith('第4')):
        print(f'[{i}] style={p.style.name} | text="{t}"')
