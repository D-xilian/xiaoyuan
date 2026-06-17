# -*- coding: utf-8 -*-
from docx import Document

doc = Document(r'd:\校园活动\基于VUE+Flask架构的校园活动发布与管理系统_论文_修改版.docx')

# 检查旧标题是否还存在
old_terms = [
    '轻量级身份认证',
    '研究问题的系统性回应',
    '学术价值与实践意义',
    '谋篇布局',
    '数据呈现',
    '文字解读',
    '关联升华',
    '四步法',
    '一.引言',
    '二．系统架构',
    '研究结果',
    '自动化通知推送',
    '被动式通知',
    '通知铃铛组件',
    'NotificationBell',
    '超过30个',
]

print('=== 旧术语残留检查 ===')
found_any = False
for term in old_terms:
    for p in doc.paragraphs:
        if term in p.text:
            print(f'  [存在] "{term}" 在段落: {p.text[:100]}...')
            found_any = True
            break

if not found_any:
    print('  ✓ 所有旧术语均已替换，无残留')
else:
    print('  ⚠ 发现部分旧术语残留')

print()
print('=== 新术语检查 ===')
new_terms = [
    '第1章 引言',
    '第2章 系统架构',
    '第3章 系统设计与实现',
    '第4章 结论',
    '3.1  功能模块划分',
    '3.1.1  核心功能模块',
    '3.1.2  用户角色与权限控制',
    '3.1.3  关键业务流程设计',
    '3.2  数据模型设计与可视化分析',
    '3.3  关键技术实现',
    '3.3.1  身份认证与权限控制',
    '3.3.2  二维码签到功能',
    '3.3.3  通知推送机制',
    '3.4  讨论与分析',
    '3.4.1  主要工作回顾',
    '3.4.2  与同类系统的比较',
    '3.4.3  系统局限与未来改进方向',
    '实践表明',
    '事件触发式通知机制',
    '已登录用户',
    '触发式生成通知',
    '30余个RESTful API',
    '较高并发访问场景',
]

all_found = True
for term in new_terms:
    found = False
    for p in doc.paragraphs:
        if term in p.text:
            found = True
            break
    if not found:
        # Also check in table
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if term in cell.text:
                        found = True
                        break
    if found:
        print(f'  ✓ "{term}"')
    else:
        print(f'  ✗ "{term}" - 未找到!')
        all_found = False

print()
if all_found:
    print('全部18个修改点应用成功! ✓')
else:
    print('存在部分未应用成功的修改点')
