# -*- coding: utf-8 -*-
from docx import Document

doc = Document(r'd:\校园活动\基于VUE+Flask架构的校园活动发布与管理系统_论文_修改版.docx')

print('=== 验证修改结果 ===')
print()

# 检查章节编号
print('--- 修改点1: 章节编号 ---')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('第') and '章' in t and len(t) < 30:
        print(f'  [{i}] {t}')

print()
print('--- 修改点2/13: 摘要 ---')
for p in doc.paragraphs:
    if '本文从系统架构设计' in p.text:
        print(f'  摘要开头: {p.text[:100]}...')
    if '实践表明' in p.text:
        print(f'  实践表明: ...{p.text[p.text.find("实践表明"):p.text.find("实践表明")+40]}...')
    if '四步法' in p.text:
        print(f'  WARNING: 仍有"四步法"残留!')

print()
print('--- 修改点3: 引言末尾 ---')
for p in doc.paragraphs:
    if '在论文结构上' in p.text:
        print(f'  {p.text}')

print()
print('--- 修改点4: 第3章小节名 ---')
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('3.') and ('功能模块' in t or '关键' in t or '数据模型' in t or '技术实现' in t or '讨论' in t):
        print(f'  {t}')
    if t.startswith('3.3.') and len(t) < 40:
        print(f'  {t}')
    if t.startswith('3.4.') and len(t) < 40:
        print(f'  {t}')

print()
print('--- 修改点5: 3.1.1通知表述 ---')
for p in doc.paragraphs:
    if '触发式生成通知' in p.text:
        print(f'  已修改: ...{p.text[p.text.find("触发式生成通知")-20:p.text.find("触发式生成通知")+20]}...')

print()
print('--- 修改点6: 3.1.3活动发布表述 ---')
for p in doc.paragraphs:
    if '提交后活动即发布上架' in p.text:
        print(f'  已修改: ...{p.text[p.text.find("提交后活动即发布上架")-20:p.text.find("提交后活动即发布上架")+20]}...')

print()
print('--- 修改点7: 3.2.1审核状态 ---')
for p in doc.paragraphs:
    if 'pending' in p.text and 'approve' in p.text:
        print(f'  已补充: ...{p.text[p.text.find("status"):p.text.find("status")+80]}...')

print()
print('--- 修改点8: API端点 ---')
for p in doc.paragraphs:
    if '30余个RESTful API端点' in p.text:
        print(f'  已修正: ...{p.text[p.text.find("30余个"):p.text.find("30余个")+60]}...')

print()
print('--- 修改点9: 表1引导语 ---')
for p in doc.paragraphs:
    if '表1列出了各功能模块的核心API接口' in p.text:
        print(f'  {p.text}')

print()
print('--- 修改点10: 3.4.2对比分析 ---')
for p in doc.paragraphs:
    if '与同类基于SSM' in p.text:
        print(f'  {p.text[:120]}...')

print()
print('--- 修改点11: 高并发表述 ---')
for p in doc.paragraphs:
    if '较高并发访问场景' in p.text:
        print(f'  已修正: ...{p.text[p.text.find("四是在较高并发"):p.text.find("四是在较高并发")+80]}...')

print()
print('--- 修改点12: 结论三段 ---')
for p in doc.paragraphs:
    if '本文设计并实现了一个基于Vue+Flask架构的校园活动发布与管理系统' in p.text:
        print(f'  段落1: {p.text[:100]}...')
    if '从实际运行效果来看' in p.text:
        print(f'  段落2: {p.text[:80]}...')
    if '受时间和条件所限' in p.text:
        print(f'  段落3: {p.text[:80]}...')

print()
print('--- 修改点14: 参考文献 ---')
for p in doc.paragraphs:
    if p.text.startswith('[7]') or p.text.startswith('[8]') or p.text.startswith('[9]'):
        print(f'  {p.text}')

print()
print('--- 修改点15: 图片引导文字 ---')
for p in doc.paragraphs:
    if '如图' in p.text and '所示' in p.text and ('饼图' in p.text or '柱状图' in p.text or '概览' in p.text):
        print(f'  引导文字: {p.text}')

print()
print('--- 修改点16: 通知机制 ---')
for p in doc.paragraphs:
    if '事件触发式通知机制' in p.text:
        print(f'  {p.text[:80]}...')

print()
print('--- 修改点17: 术语统一 ---')
# 表1
table = doc.tables[0]
for row in table.rows:
    for cell in row.cells:
        if '已登录用户' in cell.text:
            print(f'  表1: {cell.text}')
# 代码字体
for p in doc.paragraphs:
    if 'custom_login_required' in p.text:
        has_code = False
        for run in p.runs:
            if 'custom_login_required' in run.text and run.font.name == 'Courier New':
                has_code = True
        print(f'  custom_login_required 代码字体: {"正确" if has_code else "未应用"}')
# 通知提醒组件
for p in doc.paragraphs:
    if '通知提醒组件' in p.text and '定时查询' in p.text:
        print(f'  通知提醒组件已替换')

print()
print('=========== 验证完毕 ===========')
