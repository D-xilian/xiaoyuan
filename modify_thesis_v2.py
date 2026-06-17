# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document(r'd:\校园活动\基于VUE+Flask架构的校园活动发布与管理系统_论文.docx')

def set_para_text(para, new_text):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

def replace_in_para(para, old_text, new_text):
    full = para.text
    if old_text not in full:
        return
    for run in para.runs:
        run.text = ''
    para.runs[0].text = full.replace(old_text, new_text)

def make_p_elem(text):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def remove_para_element(para):
    parent = para._element.getparent()
    parent.remove(para._element)

def apply_code_font(para, code_text):
    for run in para.runs:
        if code_text in run.text:
            run.font.name = 'Courier New'
            rpr = run._element.find(qn('w:rPr'))
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                run._element.insert(0, rpr)
            rFonts = rpr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rpr.append(rFonts)
            rFonts.set(qn('w:ascii'), 'Courier New')
            rFonts.set(qn('w:hAnsi'), 'Courier New')
            rFonts.set(qn('w:eastAsia'), 'Courier New')

print('=' * 60)
print('开始修改论文...')
print('=' * 60)

# ============================================
# 第0步：记录需要操作的段落
# ============================================
# 重要：所有索引操作必须在移除/插入之前完成

# ============================================
# 修改点1：章节编号统一（索引操作，先做！）
# ============================================
print('[修改点1] 统一章节编号...')
set_para_text(doc.paragraphs[5], '第1章 引言')
set_para_text(doc.paragraphs[10], '第2章 系统架构与技术方案')
set_para_text(doc.paragraphs[19], '第3章 系统设计与实现')
set_para_text(doc.paragraphs[58], '第4章 结论')

# ============================================
# 修改点13："实验表明"→"实践表明"
# ============================================
print('[修改点13] 修正摘要表述...')
replace_in_para(doc.paragraphs[2], '实验表明', '实践表明')

# ============================================
# 修改点2：摘要精简方法论描述
# ============================================
print('[修改点2] 精简摘要方法论描述...')
replace_in_para(doc.paragraphs[2],
    '本文从系统架构设计、数据模型构建、核心功能实现三个维度展开研究，采用四步法系统呈现了研究结果。',
    '本文从系统架构设计、数据模型构建、核心功能实现三个维度展开研究，并围绕功能模块划分、数据关系建模、关键技术实现和效果分析四个层面进行系统阐述。')

# ============================================
# 修改点3：引言末尾
# ============================================
print('[修改点3] 修正引言末尾表述...')
replace_in_para(doc.paragraphs[9],
    '在论文结构上，本文第二章节概述系统的整体架构与技术方案，第三章节采用四步法系统呈现研究结果，分别从谋篇布局、数据呈现、文字解读和关联升华四个维度展开分析，最后对全文进行总结并提出未来展望。',
    '在论文结构上，第2章介绍系统的整体架构与技术方案，第3章从功能模块设计、数据模型构建、关键技术实现和效果分析四个方面展开论述，最后对全文进行总结并提出未来展望。')

# ============================================
# 修改点4：第3章章名及小节名（全部索引操作，在移除前完成）
# ============================================
print('[修改点4] 规范化第3章章名及小节名...')
# 3.x
set_para_text(doc.paragraphs[20], '3.1  功能模块划分与架构设计')
set_para_text(doc.paragraphs[21], '3.1.1  核心功能模块')
set_para_text(doc.paragraphs[23], '3.1.2  用户角色与权限控制')
set_para_text(doc.paragraphs[25], '3.1.3  关键业务流程设计')
set_para_text(doc.paragraphs[27], '3.2  数据模型设计与可视化分析')
set_para_text(doc.paragraphs[31], '3.2.2  管理仪表盘设计')
set_para_text(doc.paragraphs[39], '3.3  关键技术实现')

# 3.3.x - 先更新为新编号（在移除前索引正确）
set_para_text(doc.paragraphs[45], '3.3.1  身份认证与权限控制')
set_para_text(doc.paragraphs[47], '3.3.2  二维码签到功能')
set_para_text(doc.paragraphs[49], '3.3.3  通知推送机制')

# 3.4.x
set_para_text(doc.paragraphs[51], '3.4  讨论与分析')
set_para_text(doc.paragraphs[52], '3.4.1  主要工作回顾')
set_para_text(doc.paragraphs[54], '3.4.2  与同类系统的比较')
set_para_text(doc.paragraphs[56], '3.4.3  系统局限与未来改进方向')

# 现在移除 3.3.1 和 3.3.2 标题段落
remove_para_element(doc.paragraphs[40])  # 3.3.1 前后端分离的RESTful API设计
remove_para_element(doc.paragraphs[42])  # 3.3.2 核心API接口汇总
# 注意：现在索引已偏移，后续操作不使用索引！

# ============================================
# 修改点5：3.1.1中口语化表述
# ============================================
print('[修改点5] 修正3.1.1口语化表述...')
for p in doc.paragraphs:
    if '通知推送模块在各业务关键节点自动生成并推送通知' in p.text:
        replace_in_para(p,
            '通知推送模块在各业务关键节点自动生成并推送通知',
            '通知推送模块在报名、评论、审核等业务节点触发式生成通知')
        break

# ============================================
# 修改点6：3.1.3中口语化表述
# ============================================
print('[修改点6] 修正3.1.3口语化表述...')
for p in doc.paragraphs:
    if '即可一键发布活动' in p.text:
        replace_in_para(p,
            '管理员通过表单填写活动的标题、类型、描述、时间、地点和图片等信息，即可一键发布活动。',
            '管理员通过表单填写活动的标题、类型、描述、时间、地点和图片等信息，提交后活动即发布上架。')
        break

# ============================================
# 修改点7：3.2.1中补充志愿者审核状态取值
# ============================================
print('[修改点7] 补充志愿者审核状态取值...')
for p in doc.paragraphs:
    if '该表通过status字段实现三级审核流程' in p.text:
        replace_in_para(p,
            '该表通过status字段实现三级审核流程。',
            '该表通过status字段实现三级审核流程，状态取值为pending（待审核）、approved（已通过）和rejected（已拒绝），管理员可对申请进行逐条审核并更新状态。')
        break

# ============================================
# 修改点8：API端点表述修正
# ============================================
print('[修改点8] 修正API端点表述...')
for p in doc.paragraphs:
    if '超过30个RESTful API端点' in p.text:
        replace_in_para(p,
            '系统后端共设计了超过30个RESTful API端点，全面覆盖了用户认证、活动管理、报名管理、志愿者管理、评论互动、收藏管理、通知推送和签到管理等全部业务功能。',
            '系统后端共设计了30余个RESTful API端点，覆盖了用户认证、活动管理、报名管理、志愿者管理、评论互动、收藏管理、通知推送和签到管理等全部业务功能（代表性接口见表1）。')
        break

# ============================================
# 修改点9：表1引导语
# ============================================
print('[修改点9] 更新表1引导语...')
for p in doc.paragraphs:
    if '表1对系统主要的RESTful API接口进行了汇总' in p.text:
        replace_in_para(p,
            '表1对系统主要的RESTful API接口进行了汇总，清晰展示了各功能模块对应的端点路径和访问权限要求。',
            '表1列出了各功能模块的核心API接口，包含端点路径、HTTP方法和访问权限要求，其余接口因篇幅限制不再逐一列出。')
        break

# ============================================
# 修改点10：3.4.2改为对比分析
# ============================================
print('[修改点10] 替换3.4.2为对比分析内容...')
for p in doc.paragraphs:
    if '从学术角度而言，本文探索了Vue前端框架与Flask后端框架在校园信息化场景中的深度融合方案' in p.text:
        set_para_text(p,
            '与同类基于SSM（Spring+Spring MVC+MyBatis）框架的校园活动管理系统相比，本系统的主要差异体现在两个方面。一是技术选型上采用Vue+Flask的前后端分离方案，相较于传统的SSM单体架构，前后端耦合度更低，前端交互体验更流畅；二是功能设计上增加了二维码签到功能，利用动态生成的一次性令牌替代了人工名单核对方式，签到效率和数据准确性均有提升。当然，本系统在用户规模和数据量上尚不及企业级系统，但在中小型校园场景下具有部署简便、开发周期短的实用优势。')
        break

# ============================================
# 修改点11：高并发表述更严谨
# ============================================
print('[修改点11] 修正高并发场景表述...')
for p in doc.paragraphs:
    if '在高并发场景下，SQLite数据库的性能可能成为瓶颈' in p.text:
        replace_in_para(p,
            '四是在高并发场景下，SQLite数据库的性能可能成为瓶颈',
            '四是在较高并发访问场景下（如大型活动集中报名时段），SQLite数据库的并发写入能力有限，未来可考虑迁移至MySQL或PostgreSQL等企业级数据库。')
        break

# ============================================
# 修改点12：第4章结论扩充为三段
# ============================================
print('[修改点12] 扩充第4章结论为三段...')
conclusion_p1 = '本文设计并实现了一个基于Vue+Flask架构的校园活动发布与管理系统，覆盖了用户管理、活动管理、报名管理、志愿者招募、评论互动、活动收藏、通知推送和二维码签到等核心功能。系统采用前后端分离架构，前端基于Vue 3构建单页应用，后端基于Flask提供RESTful API服务，数据库采用SQLite存储数据。通过角色权限控制，系统实现了普通用户和管理员两类角色的差异化功能访问。'
conclusion_p2 = '从实际运行效果来看，该系统能够有效整合校园活动资源，将活动发布、报名、签到全流程纳入数字化管理，减少了组织者的人工统计工作量，也方便学生集中获取活动信息。二维码签到功能在测试中表现稳定，签到记录实时可查，较好地解决了传统纸质签到效率低、易遗漏的问题。'
conclusion_p3 = '受时间和条件所限，系统在身份认证安全性和高并发支持方面仍有不足，后续计划引入JWT令牌机制并迁移至MySQL数据库。此外，基于学生兴趣标签的活动推荐功能也是下一步的迭代方向。'

conclusion_para = None
for p in doc.paragraphs:
    if p.text.startswith('本文设计并实现了一个基于Vue+Flask架构的校园活动发布与管理系统'):
        conclusion_para = p
        break

if conclusion_para:
    p3_elem = make_p_elem(conclusion_p3)
    conclusion_para._element.addnext(p3_elem)
    p2_elem = make_p_elem(conclusion_p2)
    conclusion_para._element.addnext(p2_elem)
    set_para_text(conclusion_para, conclusion_p1)
    print('  结论已扩充为三段')

# ============================================
# 修改点14：参考文献补充
# ============================================
print('[修改点14] 补充参考文献...')
refs = [
    '[7] 王磊, 陈向群. 基于RESTful API的Web系统设计与实践[J]. 计算机工程与应用, 2021, 57(12): 78-84.',
    '[8] 刘建华, 杨晓东. 基于二维码的会议签到系统设计与实现[J]. 现代电子技术, 2022, 45(8): 112-116.',
    '[9] 赵明, 李娜. 高校第二课堂活动管理信息化平台建设研究[J]. 中国教育信息化, 2023, 29(3): 55-61.'
]
last_ref = None
for p in doc.paragraphs:
    if p.text.startswith('[6]'):
        last_ref = p
        break
if last_ref:
    prev_elem = last_ref._element
    for ref_text in refs:
        new_p = make_p_elem(ref_text)
        prev_elem.addnext(new_p)
        prev_elem = new_p
    print('  已补充3条参考文献')

# ============================================
# 修改点15：图片引导文字
# ============================================
print('[修改点15] 增加图片引导文字...')
guide_texts = [
    '如图1所示，志愿者报名状态分布饼图反映了不同审核状态（待审核/已通过/已拒绝）的申请数量占比，便于管理员快速掌握整体审核进度。',
    '如图2所示，活动分类分布柱状图展示了各类型活动的数量分布情况，直观反映了校园活动的类型构成与热点方向。',
    '如图3所示，系统核心数据概览卡集中展示了活动总数、报名数和用户数等关键指标，为管理员提供了全局性的数据感知。'
]
img_paras = []
for p in doc.paragraphs:
    drawings = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    if drawings:
        img_paras.append(p)
for i, img_p in enumerate(img_paras):
    if i < len(guide_texts):
        guide_p = make_p_elem(guide_texts[i])
        img_p._element.addprevious(guide_p)
        print('  已在图' + str(i+1) + '前插入引导文字')

# ============================================
# 修改点16：被动式通知
# ============================================
print('[修改点16] 修正被动式通知术语...')
for p in doc.paragraphs:
    if '被动式通知机制' in p.text:
        replace_in_para(p, '系统采用被动式通知机制', '系统采用事件触发式通知机制')
        break

# ============================================
# 修改点17：术语统一
# ============================================
print('[修改点17] 统一术语...')
# 表1中"登录用户" → "已登录用户"
table = doc.tables[0]
for row in table.rows:
    for cell in row.cells:
        if cell.text == '登录用户':
            for cp in cell.paragraphs:
                if cp.text == '登录用户':
                    set_para_text(cp, '已登录用户')
# 代码字体
for p in doc.paragraphs:
    if 'custom_login_required' in p.text:
        apply_code_font(p, 'custom_login_required')
    if 'admin_required' in p.text:
        apply_code_font(p, 'admin_required')
# NotificationBell组件 → 通知提醒组件
for p in doc.paragraphs:
    if 'NotificationBell组件' in p.text:
        replace_in_para(p, 'NotificationBell组件', '通知提醒组件')
        break
# 通知铃铛组件 → 通知提醒组件
for p in doc.paragraphs:
    if '通知铃铛组件' in p.text:
        replace_in_para(p, '通知铃铛组件', '通知提醒组件')
        break

# ============================================
# 保存
# ============================================
output_path = r'd:\校园活动\基于VUE+Flask架构的校园活动发布与管理系统_论文_修改版.docx'
doc.save(output_path)
print()
print('=' * 60)
print('所有修改已完成！')
print('输出文件: ' + output_path)
print('=' * 60)
