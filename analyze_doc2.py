from docx import Document

doc = Document(r'd:\校园活动\基于VUE+Flask架构的校园活动发布与管理系统_论文.docx')

# 查看表格内容
table = doc.tables[0]
print("=== 表格内容 ===")
for i, row in enumerate(table.rows):
    cells = [cell.text for cell in row.cells]
    print(f"Row {i}: {cells}")

print()

# 查看段落33, 35, 37是否包含图片（drawing元素）
for idx in [33, 34, 35, 36, 37, 38]:
    p = doc.paragraphs[idx]
    has_drawing = bool(p._element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
    has_pict = bool(p._element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'))
    # Also check for any w:drawing or v:shape
    from lxml import etree
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
             'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
             'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
             'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
             'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}
    drawings = p._element.findall('.//wp:inline', nsmap) or p._element.findall('.//wp:anchor', nsmap)
    print(f"[{idx}] text='{p.text[:80] if p.text else ''}' | has_drawing={has_drawing} | drawings={len(drawings)}")

print()

# 查看所有段落中"登录用户"出现的位置（表格中的）
for i, p in enumerate(doc.paragraphs):
    if "登录用户" in p.text:
        print(f"段落[{i}]: {p.text[:200]}")

# 查看表格单元格中的文本
print("\n=== 表格单元格详细内容 ===")
for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        if "登录用户" in cell.text:
            print(f"Table[{i}][{j}]: {cell.text}")
            # Check paragraphs in cell
            for k, cp in enumerate(cell.paragraphs):
                if "登录用户" in cp.text:
                    print(f"  CellPara[{k}]: {cp.text}")

# 检查段落[50]的"通知铃铛组件"
print(f"\n段落[50]全文: {doc.paragraphs[50].text}")
